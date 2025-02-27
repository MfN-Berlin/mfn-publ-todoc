import csv
import re
from collections import defaultdict

from docx import Document

from model import Row

# col indices
QUELLE = 0
FB1 = 1
FB2 = 2
FB3 = 3
GD = 4
DOI = 5
TITLE = 6
AUTHORSHIPS_RAW_AUTHOR_NAME = 7
MFN_AUTHORS = 8
PUBLICATION_YEAR = 9
PUBLICATION_DATE = 10
LANGUAGE = 11
ROW_TYPE = 12
JOURNAL = 13
PUBLISHER = 14
PRIMARY_LOCATION_LICENSE = 15
OPEN_ACCESS_IS_OA = 16
OPEN_ACCESS_OA_STATUS = 17
CITES_MFN_COLLECTION_SPECIMEN = 18
IS_TAXONOMIC_REVISION = 19
IS_SPECIES_DESCRIPTION = 20
ABTEILUNG_1 = 21
ABTEILUNG_2 = 22
ABTEILUNG_3 = 23
KOMMENTAR = 24
BIBLIO_VOLUME = 25
BIBLIO_ISSUE = 26
BIBLIO_FIRST_PAGE = 27
BIBLIO_LAST_PAGE = 28
LOCATIONS_LANDING_PAGE_URL = 29
TYPE_CROSSREF = 30
BOOKTITLE = 31
EDITOR = 32
PRINT = 33
EDITION = 34
BOOK_SERIE = 35

NAN = 'NaN'


def split_authors(authors):
    splits = re.split('\.,|\. &', authors[:-1])
    result = list()
    for spl in splits:
        stripped = spl.strip()
        single_whitespace = re.sub(' +', ' ', stripped)
        result.append(single_whitespace)
    return result


def add_authors(paragraph, authors, mfn_authors):
    authors_list = split_authors(authors)
    mfn_authors_list = split_authors(mfn_authors)
    for i, a in enumerate(authors_list):
        if a in mfn_authors_list:
            paragraph.add_run(f"{a}.").bold = True
        else:
            paragraph.add_run(f"{a}.")
        if i == len(authors_list) - 1:
            pass
        else:
            paragraph.add_run("; ")


def add_biblio(paragraph, first_page, last_page):
    if first_page and first_page != NAN and last_page and last_page != NAN:
        paragraph.add_run(f', {first_page}-{last_page}.')
    elif first_page and first_page != NAN:
        paragraph.add_run(f', {first_page}.')

def add_doi(paragraph, doi):
    if doi != NAN:  # print or online
        if doi.startswith('http'):
            paragraph.add_run(f' {doi}.')
        else:
            paragraph.add_run(f' DOI: https://doi.org/{doi}.')


def readCsv():
    publis = defaultdict(list)
    with open('../Master_cleaned_2024_cs_2025_02_24.CSV', newline='\n') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=';', quotechar='"')
        for row in spamreader:
            publis[row[ROW_TYPE]].append(
                Row(row_type=row[ROW_TYPE], authorships_raw_author_name=row[AUTHORSHIPS_RAW_AUTHOR_NAME],
                    mfn_authors=row[MFN_AUTHORS], publication_year=row[PUBLICATION_YEAR], title=row[TITLE],
                    journal=row[JOURNAL], biblio_volume=row[BIBLIO_VOLUME], biblio_issue=row[BIBLIO_ISSUE],
                    biblio_first_page=row[BIBLIO_FIRST_PAGE], biblio_last_page=row[BIBLIO_LAST_PAGE],
                    doi=row[DOI], open_access_is_oa=row[OPEN_ACCESS_IS_OA] == 'True', editor=row[EDITOR],
                    book_title=row[BOOKTITLE], book_serie=row[BOOK_SERIE], print=row[PRINT], publisher=row[PUBLISHER],
                    dataset='todo'))
    return publis


# Artikel  [types: peerrevartikel, wissartikel, popularartikel, beitraginpresse]
# [authorships.raw_author_name]. ([publication_year]). [title]. [journal], [biblio.volume] ([biblio.issue]), pp. [biblio.first_page]-[ biblio.last_page].
def format_article(document, articles):
    for s in articles:
        p = document.add_paragraph()
        if s.open_access_is_oa:
            p.add_run("æ")
        add_authors(p, s.authorships_raw_author_name, s.mfn_authors)
        p.add_run(f' ({s.publication_year}). {s.title}. ')
        p.add_run(s.journal).italic = True
        p.add_run(f', {s.biblio_volume}')
        if s.biblio_issue != NAN:
            p.add_run(f' ({s.biblio_issue})')
        add_biblio(paragraph=p, first_page=s.biblio_first_page, last_page=s.biblio_last_page)
        add_doi(paragraph=p, doi=s.doi)


# Monografie | Sammelband | Ausstellungskataloge [types: sammelband, buchfachbuch]
# [editor] (Ed.). ([publication_year]). [booktitle]. [book serie ]. [biblio.volume], [print ]:[ publisher]. DOI: https://doi.org/[ doi]
def format_monographie(document, articles):
    for s in articles:
        p = document.add_paragraph()
        if s.open_access_is_oa:
            p.add_run("æ")
        p.add_run(f'{s.editor} (Ed.). ')
        p.add_run(f' ({s.publication_year}). {s.book_title}. ')
        p.add_run(f' {s.book_serie}.')
        if s.biblio_volume and s.biblio_volume != NAN:
            p.add_run(f' {s.biblio_volume},')
        p.add_run(f' {s.print}:{s.publisher}.')
        add_doi(paragraph=p, doi=s.doi)

# Sammelbandbeitrag [type: sammelbandbeitrag]
# [authorships.raw_author_name ] ([publication_year]). [title]. In: [editor] (Ed.)., [booktitle], (pp. [biblio.first_page]-[ biblio.last_page]). [book serie ]. [biblio.volume], [print ]:[ publisher]. DOI: https://doi.org/[ doi]
def format_sammelbandbeitrag(document, articles):
    for s in articles:
        p = document.add_paragraph()
        if s.open_access_is_oa:
            p.add_run("æ")
        add_authors(p, s.authorships_raw_author_name, s.mfn_authors)
        p.add_run(f' ({s.publication_year}). {s.title}. ')
        if s.editor and s.editor != NAN:
            p.add_run(f'In: {s.editor} (Ed.)., ')
        p.add_run(f'{s.book_title}')
        add_biblio(paragraph=p, first_page=s.biblio_first_page, last_page=s.biblio_last_page)
        p.add_run(f' {s.book_serie}.')
        if s.biblio_volume and s.biblio_volume != NAN:
            p.add_run(f' {s.biblio_volume},')
        p.add_run(f' {s.print}:{s.publisher}.')
        add_doi(paragraph=p, doi=s.doi)


# Berichte, Arbeitspapiere, Positionpapier [Types: arbeitspapier, bericht, stellungnahme, review]
# [authorships.raw_author_name]. ([publication_year]). [title]. (pp. [biblio.first_page]-[ biblio.last_page]). [publisher]. DOI: https://doi.org/[ doi]
def format_report(document, articles):
    for s in articles:
        p = document.add_paragraph()
        if s.open_access_is_oa:
            p.add_run("æ")
        add_authors(p, s.authorships_raw_author_name, s.mfn_authors)
        p.add_run(f' ({s.publication_year}). {s.title}. ')
        add_biblio(paragraph=p, first_page=s.biblio_first_page, last_page=s.biblio_last_page)
        p.add_run(f' {s.publisher}.')
        if s.print and s.publisher and s.print != NAN and s.publisher != NAN:
            p.add_run(f' {s.print}:{s.publisher}.')
        add_doi(p, s.doi)


# Konferenzbeiträge [Types: konferenzbeitragpaper]
# [authorships.raw_author_name]. ([publication_year]). [title]. [journal], [biblio.volume] ([biblio.issue]), pp. [biblio.first_page]-[ biblio.last_page]. ]. DOI: https://doi.org/[ doi]
def format_konf(document, articles):
    format_article(document, articles)


# Datensatz [Type: datenpublikation]
# [authorships.raw_author_name]. ([publication_year]). [title]. [Dataset]. Version: ([biblio.issue]). [publisher]. DOI: https://doi.org/[ doi]
def format_datapubl(document, articles):
    for s in articles:
        p = document.add_paragraph()
        if s.open_access_is_oa:
            p.add_run("æ")
        add_authors(p, s.authorships_raw_author_name, s.mfn_authors)
        p.add_run(f' ({s.publication_year}). {s.title}. ')
        p.add_run('__DATASET__')
        p.add_run(f'. Version ({s.biblio_issue}). {s.publisher}. ')
        add_doi(paragraph=p, doi=s.doi)


def createDoc(publis):
    document = Document()

    document.add_heading('PUBLIKATIONEN', 0)
    document.add_heading('PUBLICATIONS', 0)

    # peerrevartikel
    document.add_heading('Wissenschaftliche Artikel in referierten Zeitschriften')
    document.add_heading('Scientific articles in peer-reviewed journals')

    peerrevarticle = publis.get('peerrevartikel', [])
    sorted_peerrevarticle = sorted(peerrevarticle, key=lambda r: r.authorships_raw_author_name)

    format_article(document, articles=sorted_peerrevarticle)

    # wissartikel; editorial; review
    document.add_heading('Wissenschaftliche Artikel in anderen Fachzeitschriften')
    document.add_heading('Scientific articles in other journals')

    wissartikel = publis.get('wissartikel', [])
    wissartikel.extend(publis.get('editorial', []))
    wissartikel.extend(publis.get('review', []))
    sorted_wissartikel = sorted(wissartikel, key=lambda r: r.authorships_raw_author_name)

    format_article(document, articles=sorted_wissartikel)

    # buchpopular
    document.add_heading('Populärwissenschaftliche Monografien')
    document.add_heading('Popular scientific monographs')

    buchpopular = publis.get('buchpopular', [])
    sorted_buchpopular = sorted(buchpopular, key=lambda r: r.authorships_raw_author_name)
    format_monographie(document, sorted_buchpopular)

    # sammelband
    document.add_heading('Sammelwerke – Herausgeberschaft')
    document.add_heading('Edited books – Editorship of edited volumes')

    sammelband = publis.get('sammelband', [])
    sorted_sammelband = sorted(sammelband, key=lambda r: r.authorships_raw_author_name)
    format_monographie(document, sorted_sammelband)

    # sammelbandbeitrag
    document.add_heading('Sammelbandbeiträge')
    document.add_heading('Individual contributions to edited volumes')

    sammelbandbeitrag = publis.get('sammelbandbeitrag', [])
    sorted_sammelbandbeitrag = sorted(sammelband, key=lambda r: r.authorships_raw_author_name)
    format_sammelbandbeitrag(document, sorted_sammelbandbeitrag)

    # stellungnahmen
    document.add_heading('Positionspapiere')
    document.add_heading('Position papers')

    stellungnahmen = publis.get('stellungnahmen', [])
    sorted_stellungnahmen = sorted(stellungnahmen, key=lambda r: r.authorships_raw_author_name)
    format_report(document, sorted_stellungnahmen)

    # beitraginpresse populartikel
    document.add_heading('Populärwissenschaftliche Beiträge')
    document.add_heading('Popular scientific articles')

    beitraginpresse = publis.get('beitraginpresse', [])
    beitraginpresse.extend(publis.get('populartikel', []))
    sorted_beitraginpresse = sorted(beitraginpresse, key=lambda r: r.authorships_raw_author_name)
    format_article(document, sorted_beitraginpresse)

    # konferenzbeitragpaper poster
    document.add_heading('Konferenzbeiträge')
    document.add_heading('Conference papers')

    konferenzbeitragpaper = publis.get('konferenzbeitragpaper', [])
    konferenzbeitragpaper.extend(publis.get('poster', []))
    sorted_konferenzbeitragpaper = sorted(konferenzbeitragpaper, key=lambda r: r.authorships_raw_author_name)
    format_konf(document, sorted_konferenzbeitragpaper)

    # arbeitspapier bericht project report
    document.add_heading('Berichte')
    document.add_heading('Reports')

    arbeitspapier = publis.get('arbeitspapier', [])
    arbeitspapier.extend(publis.get('bericht', []))
    arbeitspapier.extend(publis.get('project', []))
    arbeitspapier.extend(publis.get('report', []))
    sorted_arbeitspapier = sorted(arbeitspapier, key=lambda r: r.authorships_raw_author_name)
    format_report(document, sorted_arbeitspapier)

    # datenpublikation
    document.add_heading('Datenpublikationen')
    document.add_heading('Data publications')
    datenpublikation = publis.get('datenpublikation', [])
    sorted_datenpublikation = sorted(datenpublikation, key=lambda r: r.authorships_raw_author_name)

    format_datapubl(document=document, articles=sorted_datenpublikation)

    return document


publis = readCsv()
doc = createDoc(publis)
doc.save('demo2.docx')
